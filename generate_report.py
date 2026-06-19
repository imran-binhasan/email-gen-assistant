"""Generate the final assessment report as a .docx file."""

import json
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Load eval data ─────────────────────────────────────────────────────────────
with open("results/eval_report.json") as f:
    data = json.load(f)

adv  = [r for r in data if r["model"] == "gemini"]
base = [r for r in data if r["model"] == "gemini_baseline"]

SCENARIO_LABELS = {
    1:  "Follow up on job application",
    2:  "Project delivery delay notice",
    3:  "B2B SaaS outreach (casual)",
    4:  "Meeting request with CTO",
    5:  "Overdue invoice reminder",
    6:  "Constructive feedback to junior",
    7:  "Content partnership proposal",
    8:  "Critical bug escalation (urgent)",
    9:  "Scholarship inquiry",
    10: "Contract renegotiation",
}

# ── Helpers ────────────────────────────────────────────────────────────────────

def set_font(run, size=11, bold=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def add_body(doc, text, space_after=6):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    return p

def shade_row(row, hex_color="D9E1F2"):
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)

def bold_cell(cell, text, size=10, center=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.bold = True
    run.font.size = Pt(size)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def plain_cell(cell, text, size=10, center=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(size)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ── Document ───────────────────────────────────────────────────────────────────
doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(3)

# ── Cover ──────────────────────────────────────────────────────────────────────
doc.add_paragraph()
title = doc.add_heading("Email Generation Assistant", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph("AI Engineer Candidate Assessment — Final Report")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.size = Pt(13)
sub.runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x44)

meta = doc.add_paragraph("Imran Bin Hasan  ·  imranbinhasan.work@gmail.com  ·  June 2026")
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.runs[0].font.size = Pt(10)
meta.runs[0].font.color.rgb = RGBColor(0x77, 0x77, 0x77)

repo = doc.add_paragraph("GitHub: github.com/imran-binhasan/email-gen-assistant")
repo.alignment = WD_ALIGN_PARAGRAPH.CENTER
repo.runs[0].font.size = Pt(10)
repo.runs[0].font.color.rgb = RGBColor(0x1a, 0x6a, 0xc8)

doc.add_paragraph()
doc.add_page_break()

# ── Section 1: Project Overview ───────────────────────────────────────────────
add_heading(doc, "1. Project Overview", 1)
add_body(doc, (
    "This project builds a working Email Generation Assistant that takes three inputs — "
    "Intent, Key Facts, and Tone — and produces a polished professional email using a "
    "Google Gemini LLM. The assistant is evaluated with three custom metrics across "
    "10 diverse scenarios, and two distinct prompting strategies are compared."
))

add_heading(doc, "Deliverables", 2)
add_bullet(doc, "Email generator (Intent + Facts + Tone → email)")
add_bullet(doc, "3 custom evaluation metrics with LLM-as-a-Judge scoring")
add_bullet(doc, "10 test scenarios with human-written reference emails")
add_bullet(doc, "Comparative analysis: Advanced Prompt vs Zero-shot Baseline")
add_bullet(doc, "Streamlit interactive demo (streamlit_app.py)")
add_bullet(doc, "Structured output: results/eval_report.json + eval_report.csv")

doc.add_paragraph()

# ── Section 2: Prompt Engineering ────────────────────────────────────────────
add_heading(doc, "2. Prompt Engineering Techniques", 1)
add_body(doc, (
    "Three distinct techniques are layered together in the Advanced strategy. "
    "All code lives in app/prompts.py."
))

add_heading(doc, "Technique 1 — Role-Playing (System Prompt)", 2)
add_body(doc, (
    "The model is assigned a persona: \"You are an expert professional email writer with "
    "15 years of experience drafting business communications across industries.\" "
    "This anchors quality and behavioral expectations from the first token. The system "
    "prompt also lists hard rules the model must never break: no hallucinated facts, "
    "no ignored facts, no meta-commentary, always start with Subject:."
))

add_heading(doc, "Technique 2 — Few-Shot Examples", 2)
add_body(doc, (
    "Two complete, high-quality example emails are injected before every task — "
    "one formal/confident (job application follow-up) and one urgent/direct (critical bug "
    "escalation). They span opposite tone extremes, so the model has a nearby reference "
    "regardless of the requested tone. Empirically, these eliminate the tendency to "
    "default to corporate boilerplate on non-standard tones."
))

add_heading(doc, "Technique 3 — Dynamic Tone Hints", 2)
add_body(doc, (
    "A _tone_hint() function detects tone keywords at call time and injects targeted "
    "one-line guidance for the tones most prone to model drift:"
))
add_bullet(doc, "Casual: \"use contractions, breezy opener, never write 'I am writing to'\"")
add_bullet(doc, "Urgent/Direct: \"no warm opener, end with a brief status line, no 'Best regards'\"")
add_bullet(doc, "Apologetic: \"open with 'I sincerely apologize', take full responsibility first\"")
add_body(doc, (
    "This is the most impactful technique. Without it the model defaults to formal/corporate "
    "language regardless of tone — as demonstrated by the Baseline comparison."
))

add_heading(doc, "Full System Prompt", 2)
p = doc.add_paragraph()
p.style = doc.styles["No Spacing"]
run = p.add_run(
    'You are an expert professional email writer with 15 years of experience '
    'drafting business communications across industries. You write emails that are clear, '
    'human, and purposeful — never robotic or overly formal.\n\n'
    'Your emails always:\n'
    '  - Open with a greeting appropriate to the tone\n'
    '  - State the purpose clearly within the first two sentences\n'
    '  - Weave in all provided facts naturally, never as a raw bullet list\n'
    '  - Match the requested tone precisely and sustain it in every sentence\n'
    '  - Close with a sign-off that fits the tone\n'
    '  - Stay concise — every sentence earns its place\n\n'
    'Tone execution:\n'
    '  - Casual: contractions, breezy openers, short sentences, zero corporate jargon\n'
    '  - Urgent/Direct: lead with the issue, 2-3 sentences per paragraph, brief status close\n'
    '  - Enthusiastic: energetic language, varied rhythm, genuine warmth\n'
    '  - Formal/Confident: structured paragraphs, no contractions, assertive phrasing\n'
    '  - Firm but Polite: state dates/amounts plainly, no "gentle reminder"\n\n'
    'Rules you never break:\n'
    '  - Never add facts not provided in the input\n'
    '  - Never ignore a fact that was provided\n'
    '  - Never include meta-commentary before the output\n'
    '  - Always start your response directly with Subject:'
)
run.font.name = "Courier New"
run.font.size = Pt(9)
p.paragraph_format.left_indent = Inches(0.3)
p.paragraph_format.space_after = Pt(6)

doc.add_paragraph()

# ── Section 3: Custom Evaluation Metrics ──────────────────────────────────────
add_heading(doc, "3. Custom Evaluation Metrics", 1)
add_body(doc, (
    "All metrics are scored 0.0–1.0. Evaluation uses Gemini 3.1 Flash Lite as a "
    "single consistent judge (temperature=0) for all scenarios and both strategies."
))

add_heading(doc, "Metric 1 — Fact Recall Score", 2)
add_body(doc, "Definition: Fraction of input facts present (verbatim or paraphrased) in the generated email.")
add_bullet(doc, "LLM-as-a-Judge checks each fact individually with semantic matching")
add_bullet(doc, "Score = facts_present / total_facts (0.0 → 1.0)")
add_bullet(doc, "Judge returns structured JSON: per-fact verdict + overall score + justification")
add_body(doc, (
    "Why it matters: A beautiful email that omits a key deadline or discount offer is a "
    "business failure. This metric directly measures fact coverage."
))

add_heading(doc, "Metric 2 — Tone Alignment Score", 2)
add_body(doc, "Definition: How consistently the generated email sustains the requested tone.")
add_bullet(doc, "LLM-as-a-Judge with a fixed 5-point rubric (1.0 / 0.75 / 0.5 / 0.25 / 0.0)")
add_bullet(doc, "1.0 = perfect throughout; 0.75 = mostly correct with minor slip; 0.5 = partially correct")
add_bullet(doc, "Judge must provide a one-sentence justification — prevents arbitrary scoring")
add_body(doc, (
    "Why it matters: Tone mismatch is the most visible quality failure. "
    "A 'casual' email that reads like a legal document, or an 'urgent' email opening "
    "with 'I hope this message finds you well,' undermines the entire purpose."
))

add_heading(doc, "Metric 3 — Fluency Score (Hybrid)", 2)
add_body(doc, "Definition: Combined professionalism/fluency rating from LLM judgment + structural similarity to the human reference.")
add_bullet(doc, "Component A (LLM): Judge rates grammar, professionalism, and writing quality (0.0–1.0)")
add_bullet(doc, "Component B (ROUGE-L): Automated longest-common-subsequence overlap against the reference email")
add_bullet(doc, "Final score = (llm_score + rouge_l_score) / 2")
add_body(doc, (
    "Why it matters: The LLM judge alone can be lenient. ROUGE-L anchors the score to "
    "a human-written reference. The hybrid approach is methodologically richer — "
    "one automated metric and one learned judgment, averaged together."
))

doc.add_paragraph()

# ── Section 4: Evaluation Results ─────────────────────────────────────────────
add_heading(doc, "4. Evaluation Results", 1)
add_body(doc, (
    "10 scenarios × 2 strategies = 20 total evaluations. "
    "Model: Gemini 3.1 Flash Lite. Judge: Gemini 3.1 Flash Lite (temperature=0)."
))

# Per-scenario table
add_heading(doc, "Per-Scenario Scores", 2)

tbl = doc.add_table(rows=1, cols=8)
tbl.style = "Table Grid"
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

# Column widths
widths = [0.4, 2.0, 1.1, 1.1, 1.1, 1.1, 1.1, 1.1]
for i, w in enumerate(widths):
    for row in tbl.rows:
        row.cells[i].width = Inches(w)

hdr = tbl.rows[0].cells
shade_row(tbl.rows[0], "1F3864")
headers = ["#", "Scenario", "Adv FR", "Adv TA", "Adv FL", "Bas FR", "Bas TA", "Bas FL"]
for i, h in enumerate(headers):
    hdr[i].text = ""
    p = hdr[i].paragraphs[0]
    run = p.add_run(h)
    run.font.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

for a, b in zip(adv, base):
    row = tbl.add_row()
    cells = row.cells
    sid = a["scenario_id"]
    shade_row(row, "EEF2FF" if sid % 2 == 0 else "FFFFFF")

    vals = [
        str(sid),
        SCENARIO_LABELS[sid],
        str(a["fact_recall_score"]),
        str(a["tone_alignment_score"]),
        f"{a['fluency_score']:.3f}",
        str(b["fact_recall_score"]),
        str(b["tone_alignment_score"]),
        f"{b['fluency_score']:.3f}",
    ]
    for i, v in enumerate(vals):
        cells[i].text = ""
        p = cells[i].paragraphs[0]
        run = p.add_run(v)
        run.font.size = Pt(9)
        # highlight failures
        if i in (5, 6, 7) and float(v) < 1.0:
            run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
            run.font.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i != 1 else WD_ALIGN_PARAGRAPH.LEFT

doc.add_paragraph()
add_body(doc, "FR = Fact Recall · TA = Tone Alignment · FL = Fluency · Red = below 1.0 for Baseline", space_after=4)
doc.add_paragraph()

# Summary averages table
add_heading(doc, "Summary Averages", 2)

avg_tbl = doc.add_table(rows=3, cols=5)
avg_tbl.style = "Table Grid"
avg_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

avg_headers = ["Strategy", "Fact Recall", "Tone Alignment", "Fluency", "Overall"]
shade_row(avg_tbl.rows[0], "1F3864")
for i, h in enumerate(avg_headers):
    avg_tbl.rows[0].cells[i].text = ""
    p = avg_tbl.rows[0].cells[i].paragraphs[0]
    run = p.add_run(h)
    run.font.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

rows_data = [
    ("Advanced (role-play + few-shot + tone hints)", "1.000", "1.000", "0.677", "0.8925"),
    ("Baseline (zero-shot only)",                    "0.975", "0.900", "0.684", "0.853"),
]
shade_colors = ["D9EAD3", "FCE5CD"]
for ri, (label, fr, ta, fl, ov) in enumerate(rows_data):
    row = avg_tbl.rows[ri + 1]
    shade_row(row, shade_colors[ri])
    vals = [label, fr, ta, fl, ov]
    for ci, v in enumerate(vals):
        row.cells[ci].text = ""
        p = row.cells[ci].paragraphs[0]
        run = p.add_run(v)
        run.font.size = Pt(10)
        run.font.bold = (ci == 4)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT

doc.add_paragraph()

# ── Section 5: Comparative Analysis ──────────────────────────────────────────
add_heading(doc, "5. Comparative Analysis", 1)

add_heading(doc, "Which strategy performed better?", 2)
add_body(doc, (
    "The Advanced prompt strategy (0.8925 overall) outperforms the Baseline (0.853) — "
    "a +4.6% improvement in overall score. More importantly, it achieves perfect scores "
    "on both Fact Recall (1.0) and Tone Alignment (1.0) across all 10 scenarios, "
    "while the Baseline scores 0.975 and 0.900 respectively."
))

add_heading(doc, "Biggest failure mode of the Baseline", 2)
add_body(doc, (
    "Tone calibration on non-default tones. The zero-shot baseline defaults to "
    "formal/corporate language regardless of what tone is requested, because it has no "
    "reference for what 'casual' or 'urgent' looks like in practice. Three scenarios "
    "expose this clearly:"
))
add_bullet(doc, (
    "S3 (Casual, persuasive): Baseline scored 0.629 — tone=0.5, fact recall=0.75. "
    "The model wrote 'I came across your company and thought our AI-powered HR automation "
    "tool could be a great fit' — stiff, corporate phrasing — and missed one of the four "
    "provided facts. Advanced scored 0.896 — tone=1.0, fact recall=1.0."
))
add_bullet(doc, (
    "S7 (Enthusiastic, professional): Baseline tone=0.75 — enthusiasm only 'mildly conveyed.' "
    "Advanced tone=1.0."
))
add_bullet(doc, (
    "S8 (Urgent, direct): Baseline tone=0.75 — added warm opener and 'Best regards' close. "
    "Advanced tone=1.0, led directly with the problem, ended with a brief status line."
))

add_heading(doc, "Production Recommendation", 2)
add_body(doc, "Recommend the Advanced prompt strategy. Four reasons grounded in the metric data:")
add_bullet(doc, (
    "Perfect fact recall (1.0 vs 0.975): Every provided fact appears in every email. "
    "A missed deadline or discount offer in a production email is a business error."
))
add_bullet(doc, (
    "Perfect tone alignment (1.0 vs 0.9): The 10% gap is not uniform — baseline fails hard "
    "on non-standard tones. Casual and urgent tones are common in real-world use "
    "(sales outreach, incident response). A production system must handle them reliably."
))
add_bullet(doc, (
    "Comparable fluency (0.677 vs 0.684): The Baseline's slightly higher ROUGE-L is "
    "expected — it generates more 'template-like' emails that match references structurally. "
    "The Advanced prompt's lower ROUGE-L on casual scenarios is correct behaviour, "
    "not a deficiency."
))
add_bullet(doc, (
    "Reliability: Advanced prompt's consistent performance (no scenario below 0.87) vs "
    "Baseline's high variance (0.629–0.919 range) makes it far more predictable in production."
))

doc.add_paragraph()

# ── Section 6: Technical Architecture ────────────────────────────────────────
add_heading(doc, "6. Technical Architecture", 1)

add_heading(doc, "Stack", 2)
add_bullet(doc, "Model: Gemini 3.1 Flash Lite (gemini-3.1-flash-lite) — 500 RPD free tier")
add_bullet(doc, "Generation: google-genai Python SDK")
add_bullet(doc, "Evaluation judge: same model (temperature=0) for score consistency")
add_bullet(doc, "ROUGE-L: rouge-score library (automated fluency component)")
add_bullet(doc, "Retry: Tenacity — exponential backoff, 4 attempts, 10–90s window")
add_bullet(doc, "Rate limiting: 5s sleep between judge calls, 20s between scenarios (~6 RPM)")
add_bullet(doc, "Demo: Streamlit (streamlit_app.py) — Advanced vs Baseline side-by-side")

add_heading(doc, "File Structure", 2)
p = doc.add_paragraph()
run = p.add_run(
    "email-gen-assistant/\n"
    "├── app/\n"
    "│   ├── prompts.py       # All prompt templates + few-shot + tone hints\n"
    "│   ├── generator.py     # generate_mail() for both strategies\n"
    "│   └── evaluator.py     # 3 custom metrics + LLM judge calls\n"
    "├── data/\n"
    "│   ├── scenarios.json   # 10 evaluation scenarios\n"
    "│   └── references.json  # 10 human-written reference emails\n"
    "├── results/\n"
    "│   ├── eval_report.json # Full results with justifications\n"
    "│   └── eval_report.csv  # Flat table for spreadsheet review\n"
    "├── streamlit_app.py     # Interactive demo\n"
    "└── run_eval.py          # Evaluation runner"
)
run.font.name = "Courier New"
run.font.size = Pt(9)
p.paragraph_format.left_indent = Inches(0.3)

doc.add_paragraph()

# ── Footer note ───────────────────────────────────────────────────────────────
add_body(doc, (
    "Full raw evaluation data (all 20 rows, all metric scores, and LLM justifications) "
    "is available in results/eval_report.json and results/eval_report.csv in the "
    "GitHub repository."
), space_after=4)

# ── Save ──────────────────────────────────────────────────────────────────────
doc.save("Email_Generation_Assessment_Report.docx")
print("Saved: Email_Generation_Assessment_Report.docx")
